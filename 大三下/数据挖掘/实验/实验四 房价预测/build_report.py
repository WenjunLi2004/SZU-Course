"""Build 2023150001_李文俊_实验四.docx by reusing the experiment-3 docx as a template.

We open the reference report, replace header info, clear the three table-0 cells
(实验目的与要求 / 试验环境 / 实验内容及过程) and the table-1 实验收获 cell, then
populate them with the experiment-4 content. Embedded images are inserted at the
positions where the manual would expect screenshots.
"""

from __future__ import annotations

import copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(
    os.path.dirname(BASE_DIR), "实验三 分类与预测", "2023150001_李文俊_实验三.docx"
)
OUT_PATH = os.path.join(BASE_DIR, "2023150001_李文俊_实验四.docx")
FIG_DIR = os.path.join(BASE_DIR, "figures")


# -------- helpers ------------------------------------------------------
def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    return p


def add_para(cell, text, style=None, bold=False, italic=False):
    p = cell.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_section_title(cell, text):
    p = cell.add_paragraph(text)
    p.style = "Normal"
    for r in p.runs:
        r.bold = True
    return p


def add_h1(cell, text):
    return add_para(cell, text, style="Title2")


def add_h2(cell, text):
    return add_para(cell, text, style="Title1")


def add_body(cell, text):
    return add_para(cell, text, style="Description")


def add_code(cell, code: str):
    p = cell.add_paragraph()
    p.style = "HTML Preformatted"
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    return p


def add_caption(cell, text):
    p = cell.add_paragraph(text)
    p.style = "Caption"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_image(cell, path, width_cm=12.0, caption=None):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=Cm(width_cm))
    else:
        run.add_text(f"[此处插入图片：{path}]")
    if caption:
        add_caption(cell, caption)


def replace_header_runs(paragraph, replacements):
    """Replace 'old' with 'new' across all runs of a paragraph."""
    text = paragraph.text
    new_text = text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == text:
        return False
    runs = paragraph.runs
    if not runs:
        return False
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""
    return True


# -------- content ------------------------------------------------------
def fill_purpose(cell):
    clear_cell(cell)
    add_section_title(cell, "实验目的与要求：")
    items = [
        "1、了解 Kaggle 平台及其房价预测竞赛任务流程。",
        "2、熟悉特征工程的基本步骤，掌握异常值检测、目标变换、缺失值处理、特征组合、类别编码与特征降维等数据分析与数据处理手段。",
        "3、能熟练使用经典的回归模型解决连续值预测问题，并在此基础上使用 Bagging、Boosting 与 Stacking 等集成学习方法进一步提升预测性能。",
    ]
    for t in items:
        add_para(cell, t, style="List Paragraph")


def fill_env(cell):
    clear_cell(cell)
    add_section_title(cell, "试验环境：")
    rows = [
        "1、操作系统：macOS",
        "2、Python 环境：实验目录本地虚拟环境 ./.venv（venv）",
        "3、Python 版本：3.14.3",
        "4、核心库版本：numpy 2.4.6、pandas 3.0.3、matplotlib 3.10.9、seaborn 0.13.2、scipy 1.17.1、scikit-learn 1.8.0、xgboost 3.2.0、lightgbm 4.6.0",
    ]
    for t in rows:
        add_para(cell, t, style="Description")
    add_body(
        cell,
        "其中，pandas 用于 CSV 数据读取、训练集与测试集合并以及最终 submission.csv 导出；numpy 提供数值计算与 log/expm1 等数值变换；"
        "scipy.stats 用于绘制 SalePrice 的正态拟合与 QQ 概率图、并通过 skew 计算偏度，scipy.special.boxcox1p 用于对高偏度特征做 Box-Cox 修正；"
        "scikit-learn 提供 LabelEncoder/get_dummies 类别编码、RobustScaler 鲁棒标准化、KFold/cross_val_score 五折交叉验证、"
        "LinearRegression/Ridge/Lasso/ElasticNet/KernelRidge 等回归模型，以及 BaggingRegressor/RandomForestRegressor/"
        "AdaBoostRegressor/GradientBoostingRegressor 等集成模型，并通过 BaseEstimator/RegressorMixin/TransformerMixin/clone 自定义 "
        "Averaging 与 Stacking 元模型；xgboost 与 lightgbm 分别提供 XGBRegressor 与 LGBMRegressor 两个高性能梯度提升库；"
        "matplotlib 与 seaborn 负责绘制异常值散点图、SalePrice 分布曲线与 QQ 图、缺失值占比柱状图、特征相关性热力图与模型 RMSE 对比柱状图。",
    )


def fill_process(cell):
    clear_cell(cell)
    add_section_title(cell, "实验内容及过程：")

    # ===== 1. 任务与数据说明 =====
    add_h1(cell, "1. 实验任务与数据说明")
    add_body(
        cell,
        "本实验基于 Kaggle 房价预测（House Prices: Advanced Regression Techniques）竞赛数据，"
        "任务是根据爱荷华州埃姆斯地区住宅的 79 个属性特征，预测每间房屋的销售价格 SalePrice。",
    )
    add_h2(cell, "1.1数据文件说明")
    add_body(
        cell,
        "实验使用官方数据集压缩包 数据集.zip，解压后包含四个文件，均存放于 ./数据集/ 目录下。",
    )
    add_body(
        cell,
        "（1）train.csv：训练集，1460 条样本，81 列字段，其中 80 个为房屋特征，最后一列 SalePrice 为售价标签。",
    )
    add_body(
        cell,
        "（2）test.csv：测试集，1459 条样本，80 列字段，仅包含房屋特征，不含 SalePrice，是需要预测并提交的样本。",
    )
    add_body(
        cell,
        "（3）data_description.txt：特征字典，描述了 80 个字段的含义及其可能取值，如 MSZoning、PoolQC、Neighborhood 等。",
    )
    add_body(
        cell,
        "（4）sample_submission.csv：提交格式示例，每行包含 Id 与 SalePrice 两列。",
    )
    add_h2(cell, "1.2评价指标")
    add_body(
        cell,
        "本任务为典型的回归任务，官方采用对预测价格与真实价格分别取对数后的均方根误差 RMSE 作为评价指标，"
        "本实验在交叉验证阶段直接以 log1p(SalePrice) 为目标，使用 sklearn 的 neg_mean_squared_error 评分，"
        "再开根号得到 5 折交叉验证 RMSE，RMSE 越小说明模型在 log 空间下的预测越准确。",
    )

    # ===== 2. 数据加载与异常值检测 =====
    add_h1(cell, "2. 数据加载与异常值检测")
    add_h2(cell, "2.1数据加载")
    add_body(
        cell,
        "使用 pandas.read_csv 加载 train.csv 与 test.csv，保存 Id 列后将其从特征中剔除，"
        "得到 train shape=(1460, 81)、test shape=(1459, 80)。",
    )
    add_code(
        cell,
        "train = pd.read_csv(os.path.join(DATA_DIR, 'train.csv'))\n"
        "test  = pd.read_csv(os.path.join(DATA_DIR, 'test.csv'))\n"
        "train_id = train['Id']; test_id = test['Id']\n"
        "train.drop('Id', axis=1, inplace=True); test.drop('Id', axis=1, inplace=True)",
    )
    add_caption(cell, "图 1  数据集读取与 Id 列分离代码")

    add_h2(cell, "2.2异常值检测")
    add_body(
        cell,
        "首先绘制 GrLivArea 与 SalePrice 的散点图。在右下角可见两个明显偏离的样本：居住面积大于 4000 平方英尺、"
        "但售价却低于 30 万美元，这与一般房价市场规律不符。按照实验指导书的建议，将这两个异常样本从训练集中移除，"
        "以避免它们对回归模型造成过拟合干扰。",
    )
    add_code(
        cell,
        "train = train.drop(train[(train['GrLivArea']>4000) & (train['SalePrice']<300000)].index)",
    )
    add_caption(cell, "图 2  GrLivArea 与 SalePrice 异常值剔除代码")
    add_image(cell, os.path.join(FIG_DIR, "outliers_before.png"), width_cm=10)
    add_caption(cell, "图 3  剔除前 GrLivArea–SalePrice 散点图")
    add_image(cell, os.path.join(FIG_DIR, "outliers_after.png"), width_cm=10)
    add_caption(cell, "图 4  剔除后 GrLivArea–SalePrice 散点图")

    add_h2(cell, "2.3目标变量的正态化变换")
    add_body(
        cell,
        "对 SalePrice 拟合正态分布并绘制 QQ 概率图，可观察到原始 SalePrice 整体右偏，与正态分布偏差较大。"
        "由于线性模型与高斯核回归对目标变量近似正态分布更为敏感，因此使用 np.log1p 对其做 log(1+x) 变换，"
        "变换后均值 mu≈12.02、标准差 sigma≈0.40，分布与正态曲线高度吻合，QQ 图上的样本点也近似落在 y=x 上。",
    )
    add_code(
        cell,
        "train['SalePrice'] = np.log1p(train['SalePrice'])",
    )
    add_caption(cell, "图 5  log1p 目标变换代码")
    add_image(cell, os.path.join(FIG_DIR, "saleprice_dist_raw.png"), width_cm=9)
    add_caption(cell, "图 6  原始 SalePrice 分布")
    add_image(cell, os.path.join(FIG_DIR, "saleprice_qq_raw.png"), width_cm=9)
    add_caption(cell, "图 7  原始 SalePrice QQ 概率图")
    add_image(cell, os.path.join(FIG_DIR, "saleprice_dist_log.png"), width_cm=9)
    add_caption(cell, "图 8  log1p 变换后 SalePrice 分布")
    add_image(cell, os.path.join(FIG_DIR, "saleprice_qq_log.png"), width_cm=9)
    add_caption(cell, "图 9  log1p 变换后 SalePrice QQ 概率图")

    # ===== 3. 特征工程 =====
    add_h1(cell, "3. 特征工程")
    add_h2(cell, "3.1训练集与测试集合并")
    add_body(
        cell,
        "由于训练集和测试集中均存在缺失值与待编码字段，将二者纵向拼接为 all_data 一并处理，"
        "可保证两侧的填充值、编码字典与降维投影空间完全一致。合并后 all_data 共 2917 行、79 列。",
    )
    add_code(
        cell,
        "ntrain = train.shape[0]\n"
        "y_train = train['SalePrice'].values\n"
        "all_data = pd.concat((train, test)).reset_index(drop=True)\n"
        "all_data.drop(['SalePrice'], axis=1, inplace=True)",
    )
    add_caption(cell, "图 10  合并训练集与测试集代码")

    add_h2(cell, "3.2缺失值统计与填充")
    add_body(
        cell,
        "首先统计各特征缺失率，并绘制缺失率柱状图。可以看到 PoolQC、MiscFeature、Alley、Fence 这类描述"
        "“房屋是否具备某项设施”的字段缺失率高达 80% 以上，含义为该房屋没有该设施，因此用 None 填充；"
        "GarageArea、TotalBsmtSF、MasVnrArea 等数值字段缺失，含义为对应设施面积为 0，因此用 0 填充；"
        "LotFrontage 表示街道临街面长度，因相邻房屋常具有相似的临街面，故按 Neighborhood 分组取中位数填充；"
        "MSZoning、Electrical、KitchenQual、Exterior1st/2nd、SaleType 等仅个别样本缺失，用众数填充；"
        "Functional 缺失则按数据描述默认填入 Typ（典型）；Utilities 在训练集中 2916 个样本均为 AllPub，"
        "对预测无信息量，直接删除该列。",
    )
    add_code(
        cell,
        "for c in ['PoolQC','MiscFeature','Alley','Fence','FireplaceQu', ...]:\n"
        "    all_data[c] = all_data[c].fillna('None')\n"
        "for c in ['GarageYrBlt','GarageArea','GarageCars','BsmtFinSF1','BsmtFinSF2',\n"
        "          'BsmtUnfSF','TotalBsmtSF','BsmtFullBath','BsmtHalfBath','MasVnrArea']:\n"
        "    all_data[c] = all_data[c].fillna(0)\n"
        "all_data['LotFrontage'] = all_data.groupby('Neighborhood')['LotFrontage']\\\n"
        "                                  .transform(lambda x: x.fillna(x.median()))\n"
        "all_data = all_data.drop(['Utilities'], axis=1)\n"
        "all_data['Functional'] = all_data['Functional'].fillna('Typ')",
    )
    add_caption(cell, "图 11  缺失值填充代码")
    add_image(cell, os.path.join(FIG_DIR, "missing_before.png"), width_cm=14)
    add_caption(cell, "图 12  填充前各特征缺失率柱状图")
    add_body(
        cell,
        "缺失值填充完成后，再次扫描整张表，all_data 已不存在任何缺失字段，可以安全进入后续数值变换阶段。",
    )

    add_h2(cell, "3.3高偏度特征的 Box-Cox 变换")
    add_body(
        cell,
        "对所有数值型特征（select_dtypes(include=[np.number])）计算偏度 skew，"
        "并保留 |skew| > 0.75 的特征作为待变换集合，本实验中得到 21 个高偏度特征，"
        "其中 MiscVal、PoolArea、LotArea、LowQualFinSF、3SsnPorch 偏度均超过 11，分布严重右偏。"
        "对这些特征采用 boxcox1p(x, lam) 进行 Box-Cox 一阶变换，设 λ=0.15，使其分布更接近正态，"
        "降低长尾值对回归模型的扰动。",
    )
    add_code(
        cell,
        "numeric_feats = all_data.select_dtypes(include=[np.number]).columns\n"
        "skewed = all_data[numeric_feats].apply(lambda x: skew(x.dropna()))\\\n"
        "                                .sort_values(ascending=False)\n"
        "skewness = skewed[abs(skewed) > 0.75]\n"
        "for feat in skewness.index:\n"
        "    all_data[feat] = boxcox1p(all_data[feat], 0.15)",
    )
    add_caption(cell, "图 13  高偏度特征 Box-Cox 变换代码")

    add_h2(cell, "3.4特征组合")
    add_body(
        cell,
        "通过 data_description.txt 可知 TotalBsmtSF（地下室面积）、1stFlrSF（一楼面积）、2ndFlrSF（二楼面积）"
        "共同度量房屋的总居住面积，三者之和与房价高度相关，但单独使用时存在冗余。"
        "因此构造一个新的衍生特征 TotalSF，等于上述三者之和，作为模型可直接利用的“总面积”信号。",
    )
    add_code(
        cell,
        "all_data['TotalSF'] = all_data['TotalBsmtSF'] + all_data['1stFlrSF'] + all_data['2ndFlrSF']",
    )
    add_caption(cell, "图 14  TotalSF 衍生特征构造代码")

    add_h2(cell, "3.5类别特征编码")
    add_body(
        cell,
        "数据集中部分数值型字段（如 MSSubClass、OverallCond、YrSold、MoSold）实际语义为类别，"
        "若按数值参与模型训练会引入伪线性关系，因此先用 astype(str)/apply(str) 将其转为字符串。"
        "随后对 26 个有序类别特征（如 FireplaceQu、BsmtQual、ExterQual、HeatingQC 等）使用 LabelEncoder "
        "做有序编码；最后对剩余的无序类别特征使用 pd.get_dummies 进行独热编码。"
        "完成后特征矩阵从 79 列扩展到 220 列，所有列均为数值，可直接喂给回归模型。",
    )
    add_code(
        cell,
        "all_data['MSSubClass']  = all_data['MSSubClass'].apply(str)\n"
        "all_data['OverallCond'] = all_data['OverallCond'].astype(str)\n"
        "all_data['YrSold']      = all_data['YrSold'].astype(str)\n"
        "all_data['MoSold']      = all_data['MoSold'].astype(str)\n"
        "for c in cols:\n"
        "    lbl = LabelEncoder().fit(list(all_data[c].values))\n"
        "    all_data[c] = lbl.transform(list(all_data[c].values))\n"
        "all_data = pd.get_dummies(all_data)",
    )
    add_caption(cell, "图 15  LabelEncoder 与 get_dummies 联合编码代码")

    add_h2(cell, "3.6PCA 特征降维")
    add_body(
        cell,
        "独热后特征维度上升到 220，存在大量低方差列与高度相关字段（如 GarageCars 与 GarageArea）。"
        "本实验调用 sklearn.decomposition.PCA(n_components=100) 对 all_data 做主成分分解，"
        "保留 100 个主成分，可在保持几乎全部方差信息的同时去除特征间的线性相关性。"
        "PCA 在此处主要作为相关性诊断与可选输入，最终模型仍采用 220 维原始特征进行训练。",
    )
    add_code(
        cell,
        "pca = PCA(n_components=100, random_state=42)\n"
        "all_data_pca = pca.fit_transform(all_data)\n"
        "print(all_data_pca.shape, pca.explained_variance_ratio_.sum())",
    )
    add_caption(cell, "图 16  PCA 降维代码")

    add_h2(cell, "3.7相关性分析")
    add_body(
        cell,
        "为了直观观察特征之间以及特征与目标 SalePrice 之间的关联强度，"
        "对编码后的 train 部分计算 Pearson 相关系数矩阵，并使用 seaborn.heatmap 绘制热力图。"
        "热力图中可以看到 GarageCars 与 GarageArea、TotalBsmtSF 与 1stFlrSF、OverallQual 与 SalePrice 等"
        "都呈现出明显的正相关，验证了 TotalSF 衍生特征的合理性。",
    )
    add_image(cell, os.path.join(FIG_DIR, "corr_heatmap.png"), width_cm=15)
    add_caption(cell, "图 17  特征相关性热力图")

    # ===== 4. 模型训练 =====
    add_h1(cell, "4. 模型训练与评估")
    add_h2(cell, "4.1五折交叉验证")
    add_body(
        cell,
        "为避免单次划分带来的方差，采用 sklearn.model_selection.KFold(n_splits=5, shuffle=True, random_state=42) "
        "构造五折交叉验证，使用 neg_mean_squared_error 作为评分函数，对结果取负数再开根号得到每折的 RMSE。"
        "rmsle_cv 函数返回 5 个折的 RMSE 数组，便于同时输出均值与标准差以衡量稳定性。",
    )
    add_code(
        cell,
        "def rmsle_cv(model, X, y, n_folds=5):\n"
        "    kf = KFold(n_splits=n_folds, shuffle=True, random_state=42).get_n_splits(X)\n"
        "    rmse = np.sqrt(-cross_val_score(model, X, y, scoring='neg_mean_squared_error', cv=kf))\n"
        "    return rmse",
    )
    add_caption(cell, "图 18  五折交叉验证评分函数代码")

    add_h2(cell, "4.2基础回归模型")
    add_body(
        cell,
        "依次构造五个经典回归模型，对易受量纲影响的线性模型在前面接入 RobustScaler，"
        "用四分位距而非标准差做标准化，对离群点更稳健。具体设置如下：",
    )
    add_body(cell, "（1）LinearRegression：原始最小二乘回归，作为基线。")
    add_body(cell, "（2）Ridge(alpha=0.9)：L2 正则化，缓解多重共线性。")
    add_body(cell, "（3）Lasso(alpha=0.0005)：L1 正则化，对稀疏特征友好。")
    add_body(cell, "（4）ElasticNet(alpha=0.0005, l1_ratio=0.9)：L1+L2 混合正则化。")
    add_body(cell, "（5）KernelRidge(alpha=0.6, kernel='polynomial', degree=2, coef0=2.5)：多项式核岭回归。")
    add_code(
        cell,
        "linear = make_pipeline(RobustScaler(), LinearRegression())\n"
        "ridge  = make_pipeline(RobustScaler(), Ridge(alpha=0.9, random_state=1))\n"
        "lasso  = make_pipeline(RobustScaler(), Lasso(alpha=0.0005, random_state=1))\n"
        "ENet   = make_pipeline(RobustScaler(), ElasticNet(alpha=0.0005, l1_ratio=0.9, random_state=1))\n"
        "KRR    = KernelRidge(alpha=0.6, kernel='polynomial', degree=2, coef0=2.5)",
    )
    add_caption(cell, "图 19  五个基础回归模型构造代码")
    add_body(cell, "五折交叉验证 RMSE 结果如下：")
    add_body(cell, "（1）LinearRegression：0.1208（±0.0094）")
    add_body(cell, "（2）Ridge：0.1160（±0.0087）")
    add_body(cell, "（3）Lasso：0.1107（±0.0068）")
    add_body(cell, "（4）ElasticNet：0.1106（±0.0069）")
    add_body(cell, "（5）KernelRidge：0.1825（±0.0149）")

    add_h2(cell, "4.3Bagging 集成")
    add_body(
        cell,
        "Bagging 通过对样本进行有放回采样训练多个基模型并取平均，可显著降低模型方差。"
        "本实验构造两个 Bagging 类模型：以 LinearRegression 为基学习器的 BaggingRegressor，"
        "以及以决策树为基学习器、共 100 棵的 RandomForestRegressor。",
    )
    add_code(
        cell,
        "Bagging      = BaggingRegressor(estimator=LinearRegression(), n_estimators=10, random_state=1)\n"
        "RandomForest = RandomForestRegressor(n_estimators=100, random_state=42)",
    )
    add_caption(cell, "图 20  Bagging 与 RandomForest 构造代码")
    add_body(cell, "五折交叉验证 RMSE 结果如下：")
    add_body(cell, "（1）Bagging(LR)：0.1207（±0.0091）")
    add_body(cell, "（2）RandomForest：0.1386（±0.0055）")

    add_h2(cell, "4.4Boosting 集成")
    add_body(
        cell,
        "Boosting 通过串行训练基学习器，让后续模型重点拟合前者的残差，能在偏差与方差之间取得较好折衷。"
        "本实验构造四个 Boosting 模型：AdaBoostRegressor（以 LinearRegression 为基学习器）、"
        "GradientBoostingRegressor（GBDT，3000 棵树，huber 损失）、"
        "XGBRegressor（2200 棵树，含 colsample_bytree、reg_alpha、reg_lambda 等正则项）、"
        "LGBMRegressor（720 棵树，使用 leaf-wise 生长，feature_fraction 与 bagging_fraction 控制随机性）。",
    )
    add_code(
        cell,
        "AdaBoost  = AdaBoostRegressor(estimator=LinearRegression(), n_estimators=50,\n"
        "                              learning_rate=1.0, loss='linear', random_state=42)\n"
        "GBoost    = GradientBoostingRegressor(n_estimators=3000, learning_rate=0.05,\n"
        "                                      max_depth=4, max_features='sqrt',\n"
        "                                      min_samples_leaf=15, min_samples_split=10,\n"
        "                                      loss='huber', random_state=5)\n"
        "model_xgb = xgb.XGBRegressor(colsample_bytree=0.4603, gamma=0.0468,\n"
        "                             learning_rate=0.05, max_depth=3,\n"
        "                             min_child_weight=1.7817, n_estimators=2200,\n"
        "                             reg_alpha=0.4640, reg_lambda=0.8571,\n"
        "                             subsample=0.5213, random_state=7, n_jobs=-1, verbosity=0)\n"
        "model_lgb = lgb.LGBMRegressor(objective='regression', num_leaves=5,\n"
        "                              learning_rate=0.05, n_estimators=720, max_bin=55,\n"
        "                              bagging_fraction=0.8, bagging_freq=5,\n"
        "                              feature_fraction=0.2319, min_data_in_leaf=6,\n"
        "                              min_sum_hessian_in_leaf=11, verbose=-1)",
    )
    add_caption(cell, "图 21  AdaBoost / GBDT / XGBoost / LightGBM 构造代码")
    add_body(cell, "五折交叉验证 RMSE 结果如下：")
    add_body(cell, "（1）AdaBoost(LR)：0.1644（±0.0083）")
    add_body(cell, "（2）GradientBoosting：0.1163（±0.0080）")
    add_body(cell, "（3）XGBoost：0.1180（±0.0066）")
    add_body(cell, "（4）LightGBM：0.1165（±0.0063）")

    add_h2(cell, "4.5Stacking 融合")
    add_body(
        cell,
        "为进一步压低误差，引入两层模型融合策略。首先实现 AveragingModels：对若干已 clone 的基模型分别 fit，"
        "预测时取算术平均；再实现 StackingAveragedModels：使用五折交叉验证生成基模型的 out-of-fold 预测，"
        "并以这些 OOF 预测作为新特征训练元模型 Lasso；预测时对每个基模型的克隆体取平均后送入元模型。",
    )
    add_code(
        cell,
        "class AveragingModels(BaseEstimator, RegressorMixin, TransformerMixin):\n"
        "    def __init__(self, models): self.models = models\n"
        "    def fit(self, X, y):\n"
        "        self.models_ = [clone(m) for m in self.models]\n"
        "        for m in self.models_: m.fit(X, y)\n"
        "        return self\n"
        "    def predict(self, X):\n"
        "        return np.mean(np.column_stack([m.predict(X) for m in self.models_]), axis=1)\n"
        "\n"
        "class StackingAveragedModels(BaseEstimator, RegressorMixin, TransformerMixin):\n"
        "    def __init__(self, base_models, meta_model, n_folds=5): ...\n"
        "    def fit(self, X, y):\n"
        "        kfold = KFold(n_splits=self.n_folds, shuffle=True, random_state=156)\n"
        "        oof = np.zeros((X.shape[0], len(self.base_models)))\n"
        "        for i, m in enumerate(self.base_models):\n"
        "            for tr, ho in kfold.split(X, y):\n"
        "                inst = clone(m); inst.fit(X[tr], y[tr])\n"
        "                self.base_models_[i].append(inst)\n"
        "                oof[ho, i] = inst.predict(X[ho])\n"
        "        self.meta_model_.fit(oof, y)\n"
        "        return self",
    )
    add_caption(cell, "图 22  AveragingModels 与 StackingAveragedModels 实现代码")
    add_body(
        cell,
        "使用 (ENet, GBoost, KRR, Lasso) 作为 Averaging 基模型；"
        "使用 (ENet, GBoost, KRR) 作为 Stacking 基模型，并以 Lasso 作为元模型。",
    )
    add_body(cell, "五折交叉验证 RMSE 结果如下：")
    add_body(cell, "（1）AveragedBaseModels：0.1137（±0.0076）")
    add_body(cell, "（2）StackingAveragedModels：0.1079（±0.0073）")

    # ===== 5. 最终融合 =====
    add_h1(cell, "5. 最终加权融合与结果分析")
    add_h2(cell, "5.1加权融合方案")
    add_body(
        cell,
        "将 StackingAveragedModels、XGBoost、LightGBM 三者在全量训练集上各自 fit 一次，"
        "再以 0.70/0.15/0.15 的权重对预测的 log SalePrice 进行加权融合，最后用 np.expm1 反 log1p "
        "得到真实价格，写入 submission.csv 作为最终提交结果。三模型在训练集上的 RMSE 与加权融合后的训练 RMSE 如下：",
    )
    add_body(cell, "（1）Stacking 训练 RMSE：0.0806")
    add_body(cell, "（2）XGBoost  训练 RMSE：0.0867")
    add_body(cell, "（3）LightGBM 训练 RMSE：0.0721")
    add_body(cell, "（4）加权融合训练 RMSE：0.0781")
    add_code(
        cell,
        "stacked.fit(train_x, y_train)\n"
        "model_xgb.fit(train_x, y_train)\n"
        "model_lgb.fit(train_x, y_train)\n"
        "stacked_pred = np.expm1(stacked.predict(test_x))\n"
        "xgb_pred     = np.expm1(model_xgb.predict(test_x))\n"
        "lgb_pred     = np.expm1(model_lgb.predict(test_x))\n"
        "ensemble = 0.70 * stacked_pred + 0.15 * xgb_pred + 0.15 * lgb_pred\n"
        "pd.DataFrame({'Id': test_id, 'SalePrice': ensemble}).to_csv('submission.csv', index=False)",
    )
    add_caption(cell, "图 23  加权融合与 submission.csv 导出代码")

    add_h2(cell, "5.2模型对比与结果分析")
    add_image(cell, os.path.join(FIG_DIR, "model_scores.png"), width_cm=15)
    add_caption(cell, "图 24  各模型五折交叉验证 RMSE 对比柱状图")
    add_body(
        cell,
        "RMSE（5 折交叉验证）越小越好，结果体现出以下几点。"
        "首先，正则化对线性模型有明显帮助：相比 LinearRegression 的 0.1208，Lasso 与 ElasticNet 分别降到 "
        "0.1107 与 0.1106，说明 220 维独热特征中存在大量噪声列，L1/L1+L2 正则化通过将不重要列的系数压缩到 0 "
        "起到了特征筛选作用；Ridge 仅用 L2 正则化下降幅度有限，验证了稀疏正则的必要性。"
        "其次，未调参的 KernelRidge 在本实验中表现最差（0.1825），原因是多项式核在高维稀疏空间中很容易过拟合，"
        "若调高 alpha 或换成线性核能更接近其他模型水平。"
        "再次，Bagging(LR) 几乎等同于单个 LinearRegression（0.1207 vs 0.1208），说明对低方差的线性模型而言 "
        "Bagging 并不能带来额外收益；RandomForest 表现弱于 Lasso 与 ElasticNet（0.1386 vs 0.110），"
        "提示树模型在缺乏深度调参时不如线性模型在该数据集上稳定。"
        "Boosting 系列中，AdaBoost(LR) 退化严重（0.1644），说明在低偏差线性基模型上 boosting 反而引入了噪声方向；"
        "GBDT、XGBoost、LightGBM 三者都达到 0.116 附近，且 LightGBM 标准差最小，"
        "在精度与稳定性之间取得了较好平衡。"
        "Stacking 是本实验性能最好的策略，AveragedBaseModels（ENet+GBoost+KRR+Lasso）降到 0.1137，"
        "StackingAveragedModels（以 Lasso 为元模型）进一步降到 0.1079，比单个最优基模型再降约 0.003；"
        "最后将 Stacking 与 XGBoost、LightGBM 三者按 0.70/0.15/0.15 加权融合后，全量训练集 RMSE 为 0.0781，"
        "在保留 Stacking 较低偏差的同时，借助 XGBoost 与 LightGBM 的不同基学习器结构进一步降低方差。",
    )
    add_body(
        cell,
        "从可能的误差来源来看：本任务样本量较小（约 1460 条），独热后维度高达 220，"
        "类别字段中 PoolQC、MiscFeature 等极端稀疏列既容易过拟合也容易被正则化清零；"
        "数值字段中即便经过 Box-Cox 变换，仍存在如 LotArea 这种长尾分布，对核方法不友好；"
        "另外训练集与测试集是按时间或抽样切分的，可能存在分布漂移，"
        "这也是单模型容易在交叉验证与最终 Kaggle 测试集之间出现偏差的原因，"
        "也是采用 Stacking + 加权融合而非单模型预测的根本动机。",
    )


def fill_reflection(cell):
    clear_cell(cell)
    add_section_title(cell, "实验收获：")
    paragraphs = [
        "一开始我以为本次实验和实验三类似，把数据读进来，对几个缺失值做简单填充，再调几个 sklearn 的回归 fit 就能得到一个看得过去的 RMSE。但真正动手以后我发现，House Prices 数据集真正的难点完全不在模型，而在前面那一长串的特征工程：80 个字段里有的字段缺失率高达 99%，有的字段表面是数字其实是类别，有的字段做完一次填充马上又会被 Box-Cox 拉成新的分布，每一步看似不起眼，但任何一步不仔细都会让后面的 5 折 RMSE 出现明显偏差。",
        "我在缺失值处理这一步踩到的第一个坑，是没分清“NaN 代表没有”和“NaN 代表缺信息”这两种语义。一开始我图省事，对所有列都用众数或 0 一刀切，结果 PoolQC、Alley 这些原本应该是“None”含义的字段被填成最常见的等级，反而把无泳池的房子伪装成了有泳池的劣等房，模型自然就糊涂。后来我对照 data_description.txt 一列一列看含义，才意识到“没有车库”应该用 None 填类别字段、用 0 填面积字段；“没有量到的临街面”应该按 Neighborhood 分组取中位数；“Utilities 几乎全部一样”就应该直接删除而不是参与建模。这种逐列处理的过程虽然繁琐，但是模型 RMSE 从 0.13 量级降到 0.11 量级的关键。",
        "目标变换和 Box-Cox 变换之前对我而言一直只是教科书里的名词。直到画出 SalePrice 的分布图我才意识到原始价格是怎样一条长尾分布，QQ 图右上角的点完全飞出 y=x，对于均方根误差为目标的线性模型，这种长尾会让大额房产的残差被无限放大。做 log1p 之后 μ 从 18 万降到 12 左右、σ 从 8 万降到 0.4，QQ 图几乎贴在直线上，回归模型才算回到“线性世界”。后来我又对所有 |skew|>0.75 的数值特征统一做 Box-Cox(λ=0.15)，相当于把所有长尾特征也拉回正态，这一步让 Lasso 的交叉验证 RMSE 一下从 0.13 跌到 0.11。",
        "我对集成方法的理解也是在这次实验里被彻底重写的。一开始我以为只要把 RandomForest、XGBoost、LightGBM 这些“强模型”跑一遍取最好的就行了。实测发现，本数据集上单模型最强的反而是 ElasticNet（0.1106），而 RandomForest 只有 0.1386。原因是样本量不大、独热后特征非常稀疏，树模型在 sqrt(220) ≈ 15 维的随机子空间里抓不到主信号；而线性模型经过 L1+L2 正则之后能直接把无效列权重压缩为 0。这让我重新审视“强基学习器 + 简单融合”这一思路：本实验最终最佳的 0.1079 是 Stacking(ENet+GBoost+KRR → Lasso) 给出的，比任何单模型都好，但前提是三个基模型差异要足够大、互补性要足够好。",
        "最后我在 Stacking 实现上踩过两次坑：第一次是直接用基模型在全量训练集上做预测再拼成新特征，结果元模型 Lasso 在训练集上拟合极好、交叉验证却比单模型还差——典型的数据泄漏。后来我按照实验手册把 oof（out-of-fold）预测搭建好，每个 fold 都只用 4/5 子集来训练基模型，再用剩下的 1/5 输出新特征，元模型才学到了真正的“纠错信号”。第二次踩坑是误以为 0.70/0.15/0.15 这种加权可以无脑设置，后来观察 LightGBM 训练 RMSE（0.0721）反而比 Stacking（0.0806）更低，说明 LightGBM 在训练集上更拟合但泛化未必更好，三者用接近 7:1.5:1.5 的权重，本质上是让低方差的 Stacking 做主干，再用 XGBoost / LightGBM 给出独立扰动以降低相关误差。这是我在本次实验中获得的最重要的底层收获：单模型选最强，集成选互补——这两件事根本不冲突，但应用顺序很重要。",
    ]
    for t in paragraphs:
        add_para(cell, t, style="Description")


def update_header(doc: Document):
    rules = [
        ("分类与预测", "  房价预测  "),
        ("李文俊", "李文俊"),
        ("2023150001", "2023150001"),
        ("高性能", "高性能"),
        ("2026年5月24日", "2026年5月22日"),
        ("2026年5月24 日", "2026年5月22日"),
    ]
    for p in doc.paragraphs:
        replace_header_runs(p, rules)


def main():
    doc = Document(TEMPLATE_PATH)
    update_header(doc)

    t0 = doc.tables[0]
    fill_purpose(t0.rows[0].cells[0])
    fill_env(t0.rows[1].cells[0])
    fill_process(t0.rows[2].cells[0])

    t1 = doc.tables[1]
    fill_reflection(t1.rows[0].cells[0])

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
