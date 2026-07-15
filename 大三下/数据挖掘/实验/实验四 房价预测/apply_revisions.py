"""Apply two in-place fixes to 2023150001_李文俊_实验四.docx:

1. 试验环境 cell: merge the two paragraphs that accidentally split scipy in half
   (paragraph ending '...；s' + next starting 'cipy.stats ...') so scipy.stats
   becomes whole again.
2. 实验收获 cell: replace the five long paragraphs with the user-approved
   ~620-character condensed version (4 paragraphs).
"""

from __future__ import annotations

import os
from copy import deepcopy

from docx import Document

DOC = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "2023150001_李文俊_实验四.docx")


def fix_scipy_split(cell):
    paras = cell.paragraphs
    for i, p in enumerate(paras[:-1]):
        nxt = paras[i + 1]
        if p.text.rstrip().endswith("；s") and nxt.text.lstrip().startswith("cipy.stats"):
            # rewrite p: strip the orphan "s", drop the leftover "；s"
            new_p_text = p.text.rstrip()[:-2] + "；"
            # rewrite nxt: prepend "s"
            new_nxt_text = "s" + nxt.text.lstrip()
            # apply by clearing runs and writing one new run that preserves the style
            _replace_text(p, new_p_text)
            _replace_text(nxt, new_nxt_text)
            return True
    return False


def _replace_text(paragraph, new_text):
    runs = paragraph.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def replace_reflection(cell):
    paragraphs = [
        "一开始我以为这次实验和实验三类似，读数据、调几个 sklearn 的回归 fit 就能拿到不错的 RMSE。真正动手才发现，House Prices 数据集的难点根本不在模型，而在前面那段冗长的特征工程：80 个字段里有的缺失率高达 99%，有的表面是数字实际是类别，每一步不仔细都会让 RMSE 出现明显偏差。",
        "缺失值这一步我踩的坑是没分清“NaN 代表没有”和“NaN 代表缺信息”。一开始我对所有列用众数或 0 一刀切，结果 PoolQC、Alley 被填成最常见等级，把无泳池的房子伪装成了劣等房。对照 data_description.txt 一列列重新设计填充策略，RMSE 才从 0.13 量级降到 0.11 量级。目标变换同样让我意识到理论重要性：原始 SalePrice 是长尾分布，log1p 之后 σ 从 8 万降到 0.4，线性模型才算回到“线性世界”。",
        "对集成方法的理解也被这次实验重写了。我原以为把 RandomForest、XGBoost 这些“强模型”挨个跑一遍取最好的即可，实测却是 ElasticNet（0.1106）反而强过 RandomForest（0.1386）——因为样本不大、独热后特征稀疏，树模型拿不到主信号，正则化反而能直接压零无效列。最终最佳的 0.1079 来自 Stacking(ENet+GBoost+KRR → Lasso)，这证明“单模型选最强、集成选互补”才是正确顺序。",
        "Stacking 实现上我踩了两次坑：先是用全量训练集预测拼新特征，导致典型数据泄漏；改用 out-of-fold 预测后元模型才学到真正的纠错信号。再是 0.70/0.15/0.15 这种权重并不能无脑设置，本质上是让低方差的 Stacking 做主干，XGBoost 与 LightGBM 提供独立扰动以降低相关误差。这是本次实验我获得的最重要的底层收获。",
    ]

    # Strategy: keep first paragraph "实验收获：" intact, locate the existing
    # Description paragraphs and: rewrite the first N to match, then delete extras.
    existing_desc = [p for p in cell.paragraphs if p.style.name == "Description"]

    n_target = len(paragraphs)
    n_existing = len(existing_desc)

    # Rewrite the overlapping ones in place to preserve style/font
    for i in range(min(n_target, n_existing)):
        _replace_text(existing_desc[i], paragraphs[i])

    # If we have more new paragraphs than existing, append new ones cloning style
    if n_target > n_existing:
        template = existing_desc[-1] if existing_desc else None
        for i in range(n_existing, n_target):
            new_p = cell.add_paragraph()
            if template is not None:
                new_p.style = template.style
            _replace_text(new_p, paragraphs[i])

    # If existing > target, remove extras
    if n_existing > n_target:
        for p in existing_desc[n_target:]:
            p._element.getparent().remove(p._element)


def main():
    doc = Document(DOC)

    env_cell = doc.tables[0].rows[1].cells[0]
    fixed = fix_scipy_split(env_cell)
    print(f"scipy split fix applied: {fixed}")

    reflection_cell = doc.tables[1].rows[0].cells[0]
    replace_reflection(reflection_cell)
    print("Reflection replaced.")

    doc.save(DOC)
    print(f"Saved: {DOC}")


if __name__ == "__main__":
    main()
