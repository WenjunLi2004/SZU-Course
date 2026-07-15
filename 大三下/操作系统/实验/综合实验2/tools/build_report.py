from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
OUT = ROOT / "综合实验2_2023150001_李文俊.docx"


def set_run_font(run, size=11, bold=False, font_name="宋体", color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text="", size=11, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10 if level == 1 else 6, after=6, line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 12, bold=True, font_name="黑体")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=8, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=9, font_name="宋体")


def add_figure(doc, path: Path, caption: str, max_w_cm=15.4, max_h_cm=18.0):
    im = Image.open(path)
    w, h = im.size
    max_w = Cm(max_w_cm)
    max_h = Cm(max_h_cm)
    ratio = w / h
    width = max_w
    height = int(width / ratio)
    if height > max_h:
        height = max_h
        width = int(height * ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=2, line=1.0)
    p.add_run().add_picture(str(path), width=width, height=height)
    add_caption(doc, caption)


def add_code_command(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=9, font_name="Menlo")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F6FB")
    p._p.get_or_add_pPr().append(shading)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.1)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val))
    doc.add_paragraph()
    return table


def cover(doc):
    add_text(doc, "附件（四）", size=10, after=28)
    add_text(doc, "深 圳 大 学 实 验 报 告", size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=32)

    fields = [
        ("课程名称：", "操作系统"),
        ("实验项目名称：", "综合实验二 文件系统设计"),
        ("学院：", "计算机与软件学院"),
        ("专业：", "计算机科学与技术（创新班）"),
        ("指导教师：", "周明洋"),
        ("报告人：", "李文俊    学号：2023150001    班级：高性能"),
        ("实验时间：", "2026年 06 月 11 日"),
        ("实验报告提交时间：", "2026年 06 月 14 日"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=12, line=1.0)
        r1 = p.add_run(label)
        set_run_font(r1, size=12, bold=True)
        r2 = p.add_run("      " + value + "      ")
        set_run_font(r2, size=12)
    add_text(doc, "教务处制", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def report_body(doc):
    add_heading(doc, "一、实验目的与要求：", 1)
    add_text(doc, "本实验综合利用文件管理、磁盘空间管理和线程同步的相关知识，在用户态编写一个简易文件系统。程序由主线程动态申请 20M 内存，并将该连续空间抽象成空白磁盘；磁盘被划分为 1KB 大小的盘块，所有超级块、位图、文件控制块和文件内容均保存在该模拟磁盘中。")
    add_text(doc, "实验要求不设计目录结构，只管理简单文本文件；需要实现 open、rm、read、write 四类文件操作接口，并创建两个 pthread 子线程，每个线程分别创建两个文件、写入不同字符串并读回验证。")

    add_heading(doc, "二、方法、步骤：", 1)
    add_heading(doc, "1. 理论基础与设计思路", 2)
    add_text(doc, "文件系统的核心任务是把线性的存储空间组织成可命名、可分配、可回收的文件对象。本实验借鉴 EXT2 的“超级块 + 位图 + inode + 数据块”思想，但按任务要求进行简化：不实现目录和链接，只维护一个扁平的文件表。")
    add_text(doc, "超级块保存整个模拟磁盘的布局信息；inode 位图记录哪些文件控制块已经被使用；数据块位图记录 1KB 盘块的占用状态；SfsInode 记录文件名、文件大小、占用块数以及 blocks[32] 单级索引表。由于字符串长度不超过 2KB，32 个直接块已经远超实验需要，同时便于观察跨块写入。")
    add_table(doc, ["区域", "块号范围", "作用"], [
        ["SuperBlock", "Block 0", "保存 magic、块大小、总块数、位图位置、inode 表位置和空闲计数"],
        ["Inode Bitmap", "Block 1", "记录 128 个 inode 是否已分配"],
        ["Block Bitmap", "Block 2-4", "覆盖 20480 个 1KB 盘块的占用状态"],
        ["Inode Table", "Block 5-30", "保存 128 个 SfsInode 文件控制块"],
        ["Data Blocks", "Block 31 起", "保存普通文本文件内容"],
    ])
    add_figure(doc, FIG / "fig01_fs_layout.png", "图 1 简易文件系统磁盘布局结构图", max_w_cm=13.5, max_h_cm=12.5)
    add_figure(doc, FIG / "fig02_run_flow.png", "图 2 程序初始化、接口验证与双线程执行流程图", max_w_cm=6.2, max_h_cm=18.2)

    add_heading(doc, "2. 编译与运行步骤", 2)
    add_text(doc, "工程放置在 UTM 服务器的 ~/exp2_sfs 目录中，实际验证环境为 Ubuntu 20.04 aarch64，GCC 9.4.0。主要命令如下：")
    add_code_command(doc, "cd ~/exp2_sfs\nmake clean && make\n./sfs\n./run_demo.sh")

    add_heading(doc, "三、实验过程及内容：", 1)
    add_heading(doc, "1. 核心数据结构", 2)
    add_text(doc, "SuperBlock 是整个模拟磁盘的入口，记录布局和空闲资源数量；SfsInode 是简化版文件控制块，承担“文件名到数据块索引”的映射职责。")
    add_figure(doc, FIG / "fig03_structs.png", "图 3 SuperBlock 与 SfsInode 数据结构", max_w_cm=14.8, max_h_cm=12.0)

    add_heading(doc, "2. 文件系统格式化", 2)
    add_text(doc, "sfs_format 首先清空 20M 内存区，然后按固定规则写入超级块，并将元数据占用的块在 block bitmap 中标记为已使用。格式化完成后，data_block_start=31，说明 0-30 号块均属于管理信息区。")
    add_figure(doc, FIG / "fig04_format.png", "图 4 sfs_format 格式化与布局初始化代码", max_w_cm=14.8, max_h_cm=17.5)

    add_heading(doc, "3. 位图分配与回收", 2)
    add_text(doc, "inode 和数据块都通过位图管理。创建文件时分配 inode；写文件时按长度计算需要的 1KB 数据块数量并分配；删除文件时释放 inode 和所有数据块。")
    add_figure(doc, FIG / "fig05_bitmap_alloc.png", "图 5 inode 与数据块位图分配/回收代码", max_w_cm=13.2, max_h_cm=17.8)

    add_heading(doc, "4. open / write / read / rm 接口实现", 2)
    add_text(doc, "open 接口负责创建文件控制块；write 接口采用覆盖写策略，先检查空间是否足够，再释放旧块并写入新内容，避免空间不足时破坏原文件；read 接口按 blocks[] 顺序拼接内容；rm 接口释放文件占用的全部资源。")
    add_figure(doc, FIG / "fig06_write.png", "图 6 write 覆盖写入与跨块存储代码", max_w_cm=13.2, max_h_cm=17.8)
    add_figure(doc, FIG / "fig07_read_rm.png", "图 7 read 与 rm 接口代码", max_w_cm=13.0, max_h_cm=17.8)

    add_heading(doc, "5. 多线程验证", 2)
    add_text(doc, "两个子线程共享同一个 20M 模拟磁盘，因此 open、write、read、rm 和 list 均通过全局 pthread_mutex_t 包裹临界区。这样可以避免两个线程同时修改 inode 表或位图造成资源计数错误。")
    add_figure(doc, FIG / "fig08_threads.png", "图 8 pthread 子线程创建、写入、读取两个文件的核心代码", max_w_cm=14.8, max_h_cm=12.0)

    add_heading(doc, "四、实验结论：", 1)
    add_heading(doc, "1. 服务器编译结果", 2)
    add_text(doc, "在 UTM 的真实服务器环境中执行 make clean && make，GCC 未产生警告或错误，说明源代码可以在实验目标环境中直接编译。")
    add_figure(doc, FIG / "fig09_compile_server.png", "图 9 wenjun@2023150001 上的编译结果", max_w_cm=15.5, max_h_cm=6.5)

    add_heading(doc, "2. 初始化与基础接口验证", 2)
    add_text(doc, "运行 ./sfs 后，程序首先输出 20M、1KB、20480 blocks 等关键信息；随后通过 manual.txt 验证 open、write、read、覆盖写和 rm。删除后再次 read 返回 not found，说明 inode 与数据块资源已被释放。")
    add_figure(doc, FIG / "fig10_basic_server.png", "图 10 初始化、open/write/read/rm 基础接口运行结果", max_w_cm=15.5, max_h_cm=10.5)

    add_heading(doc, "3. 双线程与跨块写入验证", 2)
    add_text(doc, "两个 pthread 子线程分别创建 t1_alpha.txt、t1_beta.txt、t2_alpha.txt 和 t2_beta.txt。alpha 文件大小为 1393 bytes，占用 2 个 1KB 数据块，证明程序能够跨块写入和读回；beta 文件小于 1KB，占用 1 个数据块。最终文件表显示 4 个文件均存在，free_inodes 从 128 降到 124，free_data_blocks 从 20449 降到 20443，计数与 2+1+2+1 共 6 个数据块完全一致。")
    add_figure(doc, FIG / "fig11_threads_server.png", "图 11 双线程创建、写入、读取与最终文件表", max_w_cm=15.5, max_h_cm=13.5)

    add_heading(doc, "4. run_demo.sh 复现验证", 2)
    add_text(doc, "run_demo.sh 封装了清理、编译和运行步骤，便于现场演示时一条命令复现完整结果。截图显示脚本再次跑出相同的线程文件表和空闲资源计数。")
    add_figure(doc, FIG / "fig12_rundemo_server.png", "图 12 run_demo.sh 复现运行结果", max_w_cm=15.5, max_h_cm=13.5)

    add_heading(doc, "五、实验体会：", 1)
    add_text(doc, "本次实验把书本中较抽象的文件系统结构落到了一个可运行的用户态程序里。实现过程中最直观的体会是，文件名并不等同于文件内容，真正承载文件状态的是类似 inode 的文件控制块；文件控制块再通过块索引把逻辑文件内容映射到离散的数据块。")
    add_text(doc, "位图管理也让我更清楚地理解了文件系统为什么需要维护元数据一致性。open、write、rm 表面上是文件操作，底层实际都会改变 inode 位图、数据块位图、inode 表和超级块空闲计数。只要其中一个环节没有同步更新，最终的文件表就会出现“文件存在但块丢失”或“块已释放但仍被引用”等错误。")
    add_text(doc, "在线程部分，两个 pthread 共享同一块模拟磁盘，本质上就是多个执行流竞争同一份文件系统元数据。因此即使本实验不是跨进程文件系统，也必须用互斥锁保护临界区。通过最终的空闲 inode 和数据块计数可以验证，互斥保护后资源分配和回收保持一致。")

    add_heading(doc, "附录：工程文件说明", 1)
    add_text(doc, "sfs.c：完整简易文件系统源代码，包含磁盘布局、位图、文件接口和 pthread 验证逻辑。", after=2)
    add_text(doc, "Makefile：统一编译命令 gcc -std=c11 -Wall -Wextra -O2 -pthread sfs.c -o sfs。", after=2)
    add_text(doc, "run_demo.sh：现场演示脚本，自动清理旧二进制、重新编译并运行程序。", after=2)
    add_text(doc, "figures/*.mmd：竖版 Mermaid 流程图源码及在线渲染 URL，便于后续修改和复现。", after=2)


def teacher_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    set_paragraph_spacing(p, after=24)
    run = p.add_run("指导教师批阅意见：")
    set_run_font(run, size=14, bold=True)
    for _ in range(15):
        add_text(doc, "", after=18)
    add_text(doc, "成绩评定：", size=12, after=32)
    add_text(doc, "指导教师签字：", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, after=18)
    add_text(doc, "年     月     日", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    cover(doc)
    report_body(doc)
    teacher_page(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
