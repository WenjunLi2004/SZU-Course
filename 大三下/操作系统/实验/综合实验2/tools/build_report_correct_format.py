from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
DOCX_OUT = ROOT / "综合实验2_2023150001_李文俊_new.docx"


def set_run_font(run, size=10.5, bold=False, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold


def set_spacing(paragraph, before=0, after=3, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_border(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        elem = pg_borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            pg_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "30")
        elem.set(qn("w:color"), "000000")


def add_p(doc, text="", size=10.5, bold=False, font="宋体", align=None, before=0, after=3, line=1.25):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_spacing(p, before=before, after=after, line=line)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, font_name=font)
    return p


def add_title(doc, text, size=12, level=1):
    p = doc.add_paragraph()
    set_spacing(p, before=6 if level == 1 else 3, after=3, line=1.2)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, font_name="黑体")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=1, after=4, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=9, font_name="宋体")


def add_fig(doc, filename, caption, max_w=14.8, max_h=12.8):
    path = FIG / filename
    im = Image.open(path)
    w, h = im.size
    width = Cm(max_w)
    height = int(width / (w / h))
    if height > Cm(max_h):
        height = Cm(max_h)
        width = int(height * (w / h))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=2, after=1, line=1.0)
    p.add_run().add_picture(str(path), width=width, height=height)
    add_caption(doc, caption)


def add_code_box(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=3, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size=9, font_name="Menlo")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)


def underline_run(paragraph, label, value):
    r1 = paragraph.add_run(label)
    set_run_font(r1, size=14, bold=True, font_name="黑体")
    r2 = paragraph.add_run("  " + value + "  ")
    set_run_font(r2, size=14, bold=True, font_name="黑体")
    r2.font.underline = True


def build_cover(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    add_p(doc, "附件（四）", size=12, bold=True, font="宋体", before=18, after=70)
    add_p(doc, "深 圳 大 学 实 验 报 告", size=24, bold=True, font="黑体",
          align=WD_ALIGN_PARAGRAPH.CENTER, after=54, line=1.0)
    fields = [
        ("课程名称：", "操作系统"),
        ("实验项目名称：", "综合实验二"),
        ("学院：", "计算机与软件学院"),
        ("专业：", "计算机科学与技术（创新班）"),
        ("指导教师：", "周明洋"),
        ("报告人：", "李文俊   学号：2023150001   班级：高性能"),
        ("实验时间：", "2026 年 06 月 11 日"),
        ("实验报告提交时间：", "2026 年 06 月 14 日"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=10, after=18, line=1.0)
        underline_run(p, label, value)


def start_body_section(doc):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.95)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    add_page_border(sec)
    add_p(doc, "教务处制", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def build_body(doc):
    add_title(doc, "一、实验目的与要求：", 13, 1)
    add_p(doc, "综合利用文件管理和磁盘存储器管理的相关知识，结合对文件系统组织结构的认知，编写一个简易文件系统程序。程序在用户态由主线程动态申请 20M 内存，将这段连续空间抽象为空白磁盘，并按 1KB 盘块进行管理。")
    add_p(doc, "本实验不设计目录和链接，只考虑简单文本文件；需要在模拟磁盘上保存文件管理结构，实现 open、write、read、rm 四类接口，并用两个 pthread 子线程分别创建两个文件、写入不同字符串、再读回验证，观察文件系统元数据和数据块管理过程。")

    add_title(doc, "二、方法、步骤：（说明程序相关的算法原理或知识内容，程序设计的思路和方法，可以用流程图表述，程序主要数据结构的设计、主要函数之间的调用关系等）", 13, 1)
    add_title(doc, "1. 核心知识点", 12, 2)
    add_title(doc, "1.1 文件系统的元数据与数据区分离", 11, 2)
    add_p(doc, "真实文件系统通常把磁盘空间划分为元数据区和数据区。元数据区保存超级块、位图、inode 表等管理信息，数据区保存普通文件内容。本实验虽然只模拟一个极简文件系统，但仍保留这种核心思想：文件系统不是简单地把字符串塞进数组，而是要先维护“文件控制块到数据块”的映射关系。")
    add_title(doc, "1.2 inode 与位图管理思想", 11, 2)
    add_p(doc, "inode 的作用是描述一个文件，而文件名只是 inode 中的一个属性。本实验的 SfsInode 记录 used、name、size、block_count 和 blocks[32]。inode bitmap 表示文件控制块是否已经分配，block bitmap 表示某个 1KB 盘块是否已经占用。open 分配 inode，write 分配数据块，rm 同时释放 inode 和数据块。")
    add_title(doc, "1.3 多线程共享资源保护", 11, 2)
    add_p(doc, "两个 pthread 子线程共享同一个 20M 模拟磁盘，因此 inode 表、位图和超级块计数都属于临界资源。程序使用全局 pthread_mutex_t 包裹 open、write、read、rm、ls 等接口，保证同一时刻只有一个线程修改元数据，从而避免位图和空闲计数不一致。")

    add_title(doc, "2. 文件系统总体设计", 12, 2)
    add_p(doc, "模拟磁盘共 20M，即 20971520 字节，按 1KB 划分后共有 20480 个盘块。Block 0 保存 SuperBlock，Block 1 保存 inode bitmap，Block 2-4 保存 block bitmap，Block 5-30 保存 128 个 SfsInode，Block 31 起作为数据块区域。")
    add_fig(doc, "fig01_fs_layout.png", "图 1 简易文件系统磁盘布局结构图", max_w=14.0, max_h=9.5)
    add_fig(doc, "fig02_run_flow.png", "图 2 程序初始化、接口验证与双线程执行流程图", max_w=5.6, max_h=15.8)

    add_title(doc, "三．实验过程及内容：(对程序代码进行说明和分析，越详细越好，代码排版要整齐，可读性要高)", 13, 1)
    add_title(doc, "1. 实验编译与运行环境部署", 12, 2)
    add_p(doc, "代码放置在 UTM 服务器 ~/exp2_sfs 目录中，服务器提示符为 wenjun@2023150001，系统为 Ubuntu 20.04 aarch64，GCC 版本为 9.4.0。编译命令采用 Makefile 统一管理：")
    add_code_box(doc, "cd ~/exp2_sfs\nmake clean && make\n./sfs\n./run_demo.sh")

    add_title(doc, "2. 核心数据结构设计", 12, 2)
    add_p(doc, "SuperBlock 保存模拟磁盘的整体布局和空闲资源计数，SfsInode 保存单个文件的管理信息和直接块索引。由于实验要求字符串长度不超过 2KB，blocks[32] 的单级直接索引已经可以充分覆盖实验需要，也能展示跨块写入。")
    add_fig(doc, "fig03_structs.png", "图 3 SuperBlock 与 SfsInode 数据结构", max_w=14.4, max_h=9.4)

    add_title(doc, "3. 文件系统格式化", 12, 2)
    add_p(doc, "sfs_format 首先清空整块 20M 模拟磁盘，然后计算 inode 表占用块数，写入超级块中的各项偏移。随后把 Block 0 至 Block 30 标记为元数据占用块，因此真正的数据区从 Block 31 开始。")
    add_fig(doc, "fig04_format.png", "图 4 sfs_format 格式化与布局初始化代码", max_w=13.2, max_h=15.2)

    add_title(doc, "4. 位图分配与回收", 12, 2)
    add_p(doc, "alloc_inode_unlocked 扫描 inode bitmap，找到空闲项后置位并减少 free_inodes；alloc_block_unlocked 从 data_block_start 起扫描 block bitmap，找到空闲数据块后置位、清空块内容并减少 free_blocks。删除文件时则反向清位并恢复计数。")
    add_fig(doc, "fig05_bitmap_alloc.png", "图 5 inode 与数据块位图分配/回收代码", max_w=12.6, max_h=15.8)

    add_title(doc, "5. open、write、read、rm 接口实现", 12, 2)
    add_p(doc, "open 接口负责创建文件控制块；write 接口采用覆盖写，先检查所需数据块数量和空闲空间，再释放旧块、分配新块并按 1KB 分段写入；read 接口按 inode.blocks[] 顺序读回内容；rm 接口释放文件占用的全部数据块并回收 inode。")
    add_fig(doc, "fig06_write.png", "图 6 write 覆盖写入与跨块存储代码", max_w=12.6, max_h=15.8)
    add_fig(doc, "fig07_read_rm.png", "图 7 read 与 rm 接口代码", max_w=12.6, max_h=15.8)

    add_title(doc, "6. 双线程文件创建与读写", 12, 2)
    add_p(doc, "主线程格式化文件系统后创建两个 pthread 子线程。线程 1 负责 t1_alpha.txt 和 t1_beta.txt，线程 2 负责 t2_alpha.txt 和 t2_beta.txt。alpha 文件写入 1393 字节，占用两个 1KB 盘块；beta 文件写入 986 字节，占用一个盘块。")
    add_fig(doc, "fig08_threads.png", "图 8 pthread 子线程创建、写入、读取两个文件的核心代码", max_w=14.2, max_h=9.2)

    add_title(doc, "四、实验结论：（提供运行结果，对结果进行探讨、分析、评价，并提出结论性意见和改进想法）", 13, 1)
    add_title(doc, "1. 服务器编译验证", 12, 2)
    add_p(doc, "在 UTM 服务器上执行 make clean && make，GCC 编译过程无错误、无警告，说明源代码可以在实验指定 Linux 环境中直接构建。")
    add_fig(doc, "fig09_compile_server.png", "图 9 服务器编译验证结果", max_w=14.8, max_h=4.0)

    add_title(doc, "2. 初始化与基础接口验证", 12, 2)
    add_p(doc, "运行 ./sfs 后，程序输出 disk=20971520 bytes、block=1024 bytes、total_blocks=20480，说明 20M 模拟磁盘和 1KB 盘块划分正确。随后 manual.txt 完成 open、write、read、覆盖写、rm 全流程；rm 后再次 read 显示 not found，说明资源确已释放。")
    add_fig(doc, "fig10_basic_server.png", "图 10 初始化与基础接口运行结果", max_w=14.8, max_h=8.2)

    add_title(doc, "3. 双线程与跨块写入验证", 12, 2)
    add_p(doc, "两个子线程最终创建 4 个文件。t1_alpha.txt 和 t2_alpha.txt 均为 1393 bytes，占用 2 个数据块；t1_beta.txt 和 t2_beta.txt 均为 986 bytes，占用 1 个数据块。最终 free_inodes=124，说明 128 个 inode 中使用了 4 个；free_data_blocks=20443，说明 20449 个初始数据块中使用了 6 个，与 2+1+2+1 的占用完全一致。")
    add_fig(doc, "fig11_threads_server.png", "图 11 双线程写读与最终文件表", max_w=14.8, max_h=10.5)

    add_title(doc, "4. run_demo.sh 复现验证", 12, 2)
    add_p(doc, "run_demo.sh 将清理、编译和运行封装成一个现场演示脚本。再次运行后仍得到相同的文件表和资源计数，说明实验结果具有可复现性。")
    add_fig(doc, "fig12_rundemo_server.png", "图 12 run_demo.sh 复现运行结果", max_w=14.8, max_h=10.5)

    add_title(doc, "5. 结论性评价", 12, 2)
    add_p(doc, "实验结果表明，本程序已经按照任务书要求完成了无目录简易文件系统：主线程成功申请 20M 空间并完成格式化；文件管理结构存放在模拟磁盘内部；open、write、read、rm 均可正确工作；两个子线程能够在互斥保护下安全创建、写入和读取各自文件。")
    add_p(doc, "从资源计数看，inode 位图和数据块位图的变化与最终文件表完全对应，说明文件创建、跨块写入和删除回收逻辑是一致的。后续若继续扩展，可以在当前 inode 与位图基础上增加目录项结构，从而支持目录树、路径解析和更复杂的权限管理。")

    add_title(doc, "五、实验体会：（根据自己情况填写）", 13, 1)
    add_p(doc, "这次实验让我把课堂上“文件名、inode、数据块、位图”之间的关系真正串起来。文件名只是定位文件控制块的入口，真正描述文件状态的是 inode；而 inode 本身并不直接保存大段内容，而是通过块索引把逻辑文件映射到离散的数据块。")
    add_p(doc, "实现 write 接口时，我体会到元数据一致性的重要性。覆盖写不是简单 memcpy，而要先确认空间足够，再释放旧块并分配新块，同时更新 size、block_count、blocks[] 和空闲块计数。任何一个环节漏掉，最终都会出现文件内容、位图和超级块统计不一致的问题。")
    add_p(doc, "多线程部分则说明，即使文件系统只存在于一个进程的内存里，只要多个执行流共享同一份元数据，就必须进行同步控制。全局互斥锁虽然粒度较粗，但对于本实验的简易文件系统已经足够保证正确性。")


def build_teacher_page(doc):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)
    add_page_border(sec)
    add_p(doc, "指导教师批阅意见：", size=14, bold=True, font="黑体", before=12, after=260)
    add_p(doc, "成绩评定：", size=12, font="宋体", before=0, after=36)
    add_p(doc, "指导教师签字：", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, after=28)
    add_p(doc, "年     月     日", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)


def main():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    build_cover(doc)
    start_body_section(doc)
    build_body(doc)
    build_teacher_page(doc)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
