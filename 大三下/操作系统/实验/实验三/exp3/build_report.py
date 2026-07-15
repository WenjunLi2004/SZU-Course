from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

ROOT = Path(__file__).resolve().parent
BASE = ROOT / "template_from_word.docx"
OUT = ROOT.parent / "实验三_2023150001_李文俊_完成版_真实终端截图版.docx"
FIG = ROOT / "figures"
TERM = ROOT / "terminal_screenshots"


def next_numbering_id(numbering, tag_name):
    ids = []
    for el in numbering.findall(qn(f"w:{tag_name}")):
        key = qn("w:abstractNumId") if tag_name == "abstractNum" else qn("w:numId")
        val = el.get(key)
        if val is not None:
            ids.append(int(val))
    return max(ids, default=0) + 1


def create_paren_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "abstractNum")
    num_id = next_numbering_id(numbering, "num")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "（%1）")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "420")
    ppr.append(ind)

    for child in (start, num_fmt, lvl_text, lvl_jc, ppr):
        lvl.append(child)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(para, num_id):
    ppr = para._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)


def clear_cell(cell):
    for child in list(cell._tc):
        if child.tag.endswith("}p") or child.tag.endswith("}tbl"):
            cell._tc.remove(child)


def p(cell, text="", style="Description", align=None):
    para = cell.add_paragraph(text)
    if style:
        para.style = style
    if align is not None:
        para.alignment = align
    return para


def pnum(cell, text, num_id, style="Description"):
    para = p(cell, text, style)
    apply_numbering(para, num_id)
    return para


def code_block(cell, code):
    para = cell.add_paragraph()
    para.style = "Description"
    run = para.add_run(code.rstrip())
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
    run.font.size = Pt(8)
    return para


def add_field_run(paragraph, instr, result):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    paragraph.add_run(str(result))

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


figure_no = 0


def image(cell, filename, caption, max_width=6.2, max_height=8.6, folder=FIG):
    global figure_no
    path = folder / filename
    with Image.open(path) as im:
        w, h = im.size
    width = max_width
    height = width * h / w
    if height > max_height:
        height = max_height
        width = height * w / h

    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(width))

    figure_no += 1
    cap = cell.add_paragraph()
    cap.style = "Image"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("图 ")
    add_field_run(cap, " SEQ 图 \\* ARABIC ", figure_no)
    cap.add_run(" " + caption)


VM_LAYOUT_CODE = """static void snapshot(const char *label) {
    ++snapshot_id;
    printf("\\n===== snapshot %02d: %s =====\\n", snapshot_id, label);
    print_status_summary();
    copy_proc_file("status", label);
    copy_proc_file("maps", label);
}

for (int i = 0; i < BLOCKS; ++i) {
    snapshot("before_alloc");
    p[i] = malloc(128 * MB);
    memset(p[i], 0, 4096);
    snapshot("after_alloc");
}

int free_list[] = {2, 3, 5};
for (size_t i = 0; i < 3; ++i) {
    int idx = free_list[i] - 1;
    free(p[idx]);
    p[idx] = NULL;
    snapshot("after_free");
}"""

MAX_VMEM_CODE = """static int try_map(unsigned long long bytes) {
    void *p = mmap(NULL, bytes, PROT_NONE,
                   MAP_PRIVATE | MAP_ANONYMOUS | MAP_NORESERVE, -1, 0);
    if (p == MAP_FAILED) return 0;
    munmap(p, bytes);
    return 1;
}

while (high < cap && try_map(high)) {
    low = high;
    high <<= 1;
}
while (high - low > page) {
    unsigned long long mid = low + (high - low) / 2;
    mid -= mid % page;
    if (try_map(mid)) low = mid;
    else high = mid;
}"""

TOUCH_CODE = """unsigned char *buf = mmap(NULL, len, PROT_READ | PROT_WRITE,
                          MAP_PRIVATE | MAP_ANONYMOUS, -1, 0);
print_status_summary("after mmap before touch");

if (strcmp(mode, "read") == 0) {
    for (size_t off = 0; off < len; off += page) {
        sink ^= buf[off];
    }
} else {
    for (size_t off = 0; off < len; off += page) {
        buf[off] = (unsigned char)(buf[off] + 1);
    }
}
print_usage_delta("touch", &before);
print_status_summary("after touch");"""

LOCALITY_CODE = """static void access_global(unsigned char *buf, size_t pages, size_t page_size, int repeat) {
    for (int r = 0; r < repeat; ++r)
        for (size_t page = 0; page < pages; ++page)
            buf[page * page_size]++;
}

static void access_group16(unsigned char *buf, size_t pages, size_t page_size, int repeat) {
    const size_t group = 16;
    for (size_t base = 0; base < pages; base += group) {
        size_t end = base + group;
        if (end > pages) end = pages;
        for (int r = 0; r < repeat; ++r)
            for (size_t page = base; page < end; ++page)
                buf[page * page_size]++;
    }
}"""


def add_objectives(cell, num_id):
    p(cell, "一、实验目的与要求：", "Normal")
    pnum(cell, "加深对内存分配与使用操作的直观认识；", num_id)
    pnum(cell, "掌握Linux操作系统的内存分配与释放接口，理解malloc、free、mmap、munmap与内核虚拟内存区域之间的关系；", num_id)
    pnum(cell, "学会结合/proc/pid/maps、/proc/pid/status和/proc/pid/smaps观测进程逻辑地址、虚拟空间和物理页占用的变化；", num_id)
    pnum(cell, "通过读触页、写触页、双进程竞争和访问局部性对比，验证按需分页、页面置换与工作集局部性对系统性能的影响。", num_id)


def add_methods(cell, constraint_num_id):
    p(cell, "二、方法、步骤：(说明程序相关的算法原理或知识内容，程序设计的思路和方法，可以用流程图表述，程序主要数据结构的设计、主要函数之间的调用关系等)", "Normal")

    p(cell, "1. 核心知识点", "Title2")
    p(cell, "1.1 进程虚拟地址空间与VMA组织", "Title1")
    p(cell, "Linux为每个进程维护独立的用户态虚拟地址空间，进程看到的指针值本质上是虚拟地址。内核并不是按“整块进程空间”粗糙管理，而是用一组虚拟内存区域描述代码段、数据段、堆、匿名映射区、动态库、栈、vvar和vdso等区间。/proc/pid/maps正是这些VMA的线性化展示：每一行给出起止地址、读写执行权限、偏移、设备号、inode和映射文件。实验中六个128MB块能合并成一个768MB匿名映射，释放后又能切成多个128MB映射段，直接说明VMA会随用户态分配释放发生拆分与合并。")
    p(cell, "1.2 malloc、free、mmap与munmap的分工", "Title1")
    p(cell, "malloc只是C库接口，真正改变进程地址空间的是brk或mmap这类系统调用。小块内存通常从堆顶扩展或从已有堆块中切分，大块内存则更常由glibc通过匿名mmap向内核申请独立VMA；free释放大块时，glibc可以调用munmap把整段VMA还给内核。本实验连续申请128MB和1024MB，规模远大于普通堆小块，因此maps中观察到的主要是匿名rw-p映射而不是传统heap持续变长。")
    p(cell, "1.3 /proc状态字段与smaps解释", "Title1")
    p(cell, "/proc/pid/status提供进程级摘要：VmSize表示进程已经拥有的虚拟地址空间规模，VmData主要反映数据段、堆和匿名映射，VmRSS表示当前驻留在物理内存中的页面总量，RssAnon表示匿名页驻留量，VmSwap表示已经换出到交换区的匿名页。/proc/pid/smaps则把这些统计拆到每个VMA之下，可以进一步看到Size、Rss、Pss、Shared_Clean、Private_Dirty等字段。前者适合做整体对比，后者适合定位某一段映射是否真正占用了物理页。")
    p(cell, "1.4 页表、4KB页与缺页异常", "Title1")
    p(cell, "虚拟地址并不直接访问物理内存，CPU需要通过页表把虚拟页号翻译成物理页框号。实验中访问步长固定为sysconf返回的页大小，即4KB，是为了让每次循环都跨入一个新的虚拟页。当页表项不存在或权限不满足时，处理器进入内核缺页异常处理流程；内核根据VMA权限、访问类型和内存压力决定是映射共享零页、分配新物理页、从swap换入，还是向进程发送异常信号。")
    p(cell, "1.5 按需分页、共享零页与写时分配", "Title1")
    p(cell, "mmap或malloc成功返回只说明虚拟区间已经建立，并不说明物理页已经立即分配。匿名页首次读访问时，内核可以把大量虚拟页临时映射到同一个只读零页，因此VmSize迅速增长而RssAnon几乎不变；首次写访问时，内核必须为该页分配私有物理页并更新页表权限，所以RssAnon会随着写过的页数上升。读触页和写触页的对比，是理解“虚拟空间申请”和“物理页兑现”区别的关键。")
    p(cell, "1.6 overcommit、RLIMIT_AS与最大虚拟映射", "Title1")
    p(cell, "64位进程的虚拟地址空间远大于物理内存，Linux还允许在一定策略下进行内存超量承诺。max_vmem.c使用PROT_NONE和MAP_NORESERVE测试单次连续匿名映射，是为了只测地址空间能否保留，而不强迫内核提前准备对应物理页。若RLIMIT_AS没有限制，实验测得的上界主要由用户态地址空间布局、体系结构地址宽度和内核映射策略决定，因此会远大于3.8GB物理内存。")
    p(cell, "1.7 页面置换、swap与物理页竞争", "Title1")
    p(cell, "当所有活跃进程的匿名页工作集超过可用物理内存时，内核会在内存回收路径中选择一部分不活跃匿名页写入swap，并回收其物理页框供其他进程使用。双进程竞争实验中，A进程先写入并保持1792MB，B进程再写入4096MB，A的VmRSS被大幅压低而VmSwap升高，说明它的虚拟地址仍然存在，但对应物理页已经被换出。")
    p(cell, "1.8 访问局部性与工作集大小", "Title1")
    p(cell, "访问总量相同并不代表性能相同。全空间逐页重复扫描会让下一次访问某页时距离上次访问太久，页面很可能已经被换出；16页分组访问则把短时间工作集控制在64KB左右，组内重复访问发生在页面仍驻留时，因此主缺页次数和运行时间会显著下降。这一现象把课本中的局部性原理具体落实到了主缺页次数和运行时间两个可测指标上。")

    p(cell, "2. 实验设计思路", "Title2")
    p(cell, "2.1 程序模块划分", "Title1")
    p(cell, "本次实验采用四个独立程序完成观测，避免把所有现象揉在一个难以解释的大程序中。vm_layout.c负责连续申请六个128MB块、释放2号、3号、5号块、再申请1024MB和64MB，并在每一步前后复制/proc/pid/maps与/proc/pid/status。max_vmem.c使用MAP_NORESERVE和PROT_NONE执行单次连续匿名映射的指数扩张与二分测试，用于测量虚拟地址空间边界。touch_pages.c统一完成读触页、写触页和保持映射三种模式，用于对照VmRSS与VmSwap变化。locality_compare.c固定总访问次数，分别执行全空间逐页循环和16页分组访问，用getrusage记录主缺页次数和运行时间。")
    p(cell, "2.2 关键约束与防错机制", "Title1")
    pnum(cell, "所有程序都只在Linux虚拟机中运行，因为macOS没有Linux风格的/proc/pid/maps和/proc/pid/status，若在宿主机运行会失去实验对象。", constraint_num_id)
    pnum(cell, "大块虚拟内存测试使用PROT_NONE和MAP_NORESERVE，目的是测虚拟地址可映射边界，而不是无意义地把物理内存写爆。", constraint_num_id)
    pnum(cell, "读写触页实验把页面步长固定为sysconf返回的4KB，使每一次访问都对应一个明确的页级事件，避免把缓存行或字节级循环误当成页分配证据。", constraint_num_id)
    pnum(cell, "双进程竞争实验不依赖A进程自身日志，而是在B进程运行前后由外部读取/proc/A/status，避免stdout重定向的全缓冲影响实时观察。", constraint_num_id)

    p(cell, "3. 流程图表示", "Title2")
    p(cell, "本实验的核心生命周期采用竖向Mermaid思路组织，源码形式为graph TD，节点均以双引号包裹，实际报告中以图片形式呈现。")
    p(cell, "graph TD\nA[\"创建exp3工作目录\"] --> B[\"编译四个实验程序\"]\nB --> C[\"观察虚存分配与空洞复用\"]\nC --> D[\"测量最大虚拟映射\"]\nD --> E[\"执行读触页与写触页对比\"]\nE --> F[\"运行双进程物理页竞争\"]\nF --> G[\"比较全空间循环与16页分组访问\"]")
    image(cell, "flowchart.png", "内存管理实验整体流程图", max_width=5.8)


def add_process(cell, compare_num_id):
    p(cell, "三．实验过程及内容：(对程序代码进行说明和分析，越详细越好，代码排版要整齐，可读性要高)", "Normal")

    p(cell, "1. 实验编译与运行环境部署", "Title2")
    p(cell, "1.1 工作目录与编译命令", "Title1")
    p(cell, "宿主机和Ubuntu虚拟机中均建立exp3目录，虚拟机实际工作目录为/home/wenjun/exp3。编译命令如下：")
    p(cell, "gcc -Wall -Wextra -O0 -g -o vm_layout vm_layout.c")
    p(cell, "gcc -Wall -Wextra -O2 -g -o max_vmem max_vmem.c")
    p(cell, "gcc -Wall -Wextra -O2 -g -o touch_pages touch_pages.c")
    p(cell, "gcc -Wall -Wextra -O2 -g -o locality_compare locality_compare.c")
    p(cell, "其中vm_layout关闭优化是为了让观测顺序和源码顺序保持直观，其余三个程序开启O2是为了减少循环本身的额外开销，使耗时差异更集中地反映内存系统行为。")
    image(cell, "01_environment.png", "Ubuntu虚拟机实验环境与工作目录终端页面", max_width=6.2, folder=TERM)

    p(cell, "2. 虚拟内存布局与空洞复用实验", "Title2")
    p(cell, "2.1 快照记录与分配释放代码", "Title1")
    p(cell, "vm_layout.c的核心思想是在每个分配与释放动作前后调用snapshot。snapshot内部读取/proc/getpid()/status中的VmSize、VmRSS、VmData等字段，并把完整maps复制到snapshots目录。这样每个动作都有可回溯的状态文件，而不是只依赖终端瞬时输出。")
    p(cell, "核心代码如下：")
    code_block(cell, VM_LAYOUT_CODE)
    p(cell, "这段代码的关键不在malloc本身，而在snapshot插入点。snapshot在每次动作前后读取同一个进程的status和maps，保证VmSize、VmRSS、VmData与VMA地址变化能够一一对应。p数组保留六个块的首地址，使后续free能够精确释放2号、3号、5号；每次malloc后只写首个4KB页，是为了触发最小限度的页表建立，同时避免把768MB空间全部兑现为物理页。free_list中的编号先减1再访问数组，保证实验报告中的1号到6号与C语言数组下标之间不会错位。")

    p(cell, "2.2 maps片段与64MB落点分析", "Title1")
    image(cell, "03_maps_snapshots.png", "终端查看maps关键匿名映射片段", max_width=6.2, folder=TERM)
    p(cell, "从maps片段可以看到，六个128MB块最初被合并为一个768MB匿名VMA。释放2号、3号、5号后，VMA被切分为三个仍被占用的128MB段，中间形成128MB和256MB两类空洞。1024MB申请无法放入这些空洞，因而被安排到更低地址区域，并与低地址端的6号块合并为1152MB映射。64MB申请实测地址为0xffff7f4be010，落在2号与3号合并空洞的高地址端，随后与1号块合并为192MB映射。这个结果说明我的预测方向正确：小块会优先复用已有空洞，但具体落点遵循内核自高向低查找空闲区间的策略。")

    p(cell, "3. 最大虚拟内存空间测试", "Title2")
    p(cell, "3.1 指数扩张与二分搜索代码", "Title1")
    p(cell, "max_vmem.c没有触碰映射页，而是使用PROT_NONE与MAP_NORESERVE创建纯虚拟保留区。程序先从1GB开始指数翻倍寻找失败上界，再在成功下界与失败上界之间按页大小二分，从而得到单个进程一次连续匿名映射能够成功建立的最大长度。")
    p(cell, "核心代码如下：")
    code_block(cell, MAX_VMEM_CODE)
    p(cell, "try_map只负责验证一个长度是否能建立连续匿名映射，成功后立刻munmap释放，避免测试过程本身消耗大量地址空间。PROT_NONE让该区域不可读写，MAP_NORESERVE避免把虚拟保留误转化为物理承诺。指数扩张阶段快速定位失败上界，二分阶段再按页大小收敛到最大可映射长度，这样既能避免线性试探太慢，也能保证结果粒度与页表管理单位一致。")

    p(cell, "4. 读触页与写触页对比实验", "Title2")
    p(cell, "4.1 触页程序核心代码", "Title1")
    p(cell, "touch_pages.c把映射大小、访问模式和保持时间都做成命令行参数。读模式对每个页面执行buf[off]读取，写模式对每个页面执行buf[off]=buf[off]+1。两种模式的虚拟映射长度完全相同，差异只来自读访问和写访问触发的内核页分配路径。")
    p(cell, "核心代码如下：")
    code_block(cell, TOUCH_CODE)
    p(cell, "mmap之后立即打印status，是为了证明“虚拟区间已建立但尚未触页”时VmSize和VmRSS会分离。循环变量off每次增加page，因此每次访问都落到新页的页首。读分支只把数据异或到sink，防止编译器把读循环优化掉；写分支执行原地加1，使每个匿名页必须获得私有物理页。循环结束后再次读取status和getrusage，就能把次缺页、主缺页、RssAnon和VmSwap联系起来分析。")

    p(cell, "5. 提高部分实验实现", "Title2")
    p(cell, "5.1 双进程物理页竞争设计", "Title1")
    p(cell, "提高部分第一项使用touch_pages的write-hold模式完成。A进程先写入1792MB并保持180秒，外部进程在B运行前读取/proc/A/status；随后B进程写入4096MB，结束后再次读取/proc/A/status。由于B单独写入容量已经达到虚拟机物理内存规模，内核必须在A和B的匿名页之间做回收选择。")
    p(cell, "5.2 16页分组访问代码", "Title1")
    p(cell, "locality_compare.c设置两个访问函数。global模式按“全空间扫一轮、再扫下一轮”的顺序访问，group16模式则每次取16个连续页，把这些页的重复访问全部完成后再进入下一组。两者总页访问次数完全一致，区别只在工作集是否具有短时间局部性。")
    p(cell, "核心代码如下：")
    code_block(cell, LOCALITY_CODE)
    p(cell, "access_global把repeat放在最外层，意味着第1轮访问完4GB后才会再次回到第1页，此时旧页很可能已被换出。access_group16把16个连续页作为一个小工作集，在进入下一组前完成全部重复访问，短时间内反复命中的页面只有64KB左右。两个函数都执行buf[page * page_size]++，因此访问总次数和写入粒度相同，最终时间差异主要来自页面置换压力而不是代码工作量差异。")

    p(cell, "6. 机制横向对比总结", "Title2")
    pnum(cell, "从数据介质看，VmSize描述的是虚拟地址区间，VmRSS描述的是当前驻留在物理内存中的页，VmSwap描述的是被换出到交换区的匿名页。单看malloc返回值无法证明物理内存是否真的被占用，必须把这三个字段合起来看。", compare_num_id)
    pnum(cell, "从触发路径看，读匿名页主要经过缺页异常映射映射零页，写匿名页则必须分配私有物理页并建立可写页表项，因此写触页才会让RssAnon快速升高。", compare_num_id)
    pnum(cell, "从调度敏感度看，单进程虚拟映射只受地址空间和VMA策略影响，双进程密集写入则会触发页面回收和swap；访问模式虽然不改变访问总量，却会改变页面在再次访问前是否仍留在内存中。", compare_num_id)


def add_results(cell):
    p(cell, "四、实验结论：（提供运行结果，对结果进行探讨、分析、评价，并提出结论性意见和改进想法）", "Normal")

    p(cell, "1. 运行结果验证", "Title2")
    p(cell, "本节运行证据均为在浅色iTerm2终端中通过SSH进入Ubuntu虚拟机后截取的真实终端页面。对于需要长时间运行或产生明显换页压力的实验，终端直接读取程序此前运行时保存的logs原始输出，保留实际PID、地址、缺页数和swap数据，避免为截图重复制造内存压力。")
    p(cell, "1.1 虚存分配、释放与64MB预测验证", "Title1")
    image(cell, "02_vm_layout_output.png", "终端显示虚存分配实验原始日志", max_width=6.2, folder=TERM)
    p(cell, "运行结果显示，六次128MB申请后VmSize从1924KB增长到788380KB，释放2号、3号、5号后回落到395152KB。1024MB申请后VmSize升至1443732KB，64MB申请后升至1509272KB。64MB实测地址0xffff7f4be010位于原2号和3号释放后形成的高地址空洞内，验证了实验前关于“优先复用已有足够大空洞”的判断。")
    p(cell, "用户进程空间分配在虚拟地址视角上属于离散分配。每个VMA可以连续，但多个VMA之间可以被库、栈、vvar、vdso和已释放空洞隔开。就本次mmap大块分配行为看，内核更接近自高地址向低地址扫描的首次适应，而不是为了找到最小剩余空间而全局搜索的最佳适应。用户空间存在外部碎片，但虚拟地址空间很大，且页表映射把虚拟连续性和物理连续性解耦，所以这种碎片通常不会像传统连续物理分配那样立即导致严重失败。")

    p(cell, "1.2 单进程最大虚拟空间验证", "Title1")
    image(cell, "04_max_vmem_output.png", "终端显示最大虚拟映射测试原始输出", max_width=6.2, folder=TERM)
    p(cell, "测试机RLIMIT_AS显示为unlimited，程序从1GB一直指数扩张到128TiB仍能成功，首次失败上界为256TiB，最终二分得到最大单次连续匿名虚拟映射约174763.53GiB，即170.67TiB。该值远大于物理内存3.8GiB，说明这里测到的是64位进程用户地址空间的可映射边界，而不是实际DRAM容量。")

    p(cell, "1.3 读触页与写触页验证", "Title1")
    image(cell, "05_touch_compare_output.png", "终端显示3GB读触页与写触页状态对比", max_width=6.2, folder=TERM)
    p(cell, "3GB读触页后VmSize保持3147652KB，但VmRSS仅约1212KB，RssAnon约80KB，VmSwap为0；这说明读操作虽然触发了786434次次缺页，但基本没有分配私有物理页。3GB写触页后VmRSS升至2750784KB，RssAnon升至2749572KB，同时VmSwap达到396160KB，证明写访问才会把虚拟页真正落实为匿名物理页，并在物理内存不足时推动一部分页进入交换区。")

    p(cell, "1.4 双进程竞争验证", "Title1")
    image(cell, "06_competition_output.png", "终端显示两个进程竞争物理页运行结果", max_width=6.0, folder=TERM)
    p(cell, "A进程写入并保持1792MB后，运行B进程前A的VmRSS为1836048KB，VmSwap为0。B进程写入4096MB后，A的VmRSS被压缩到1064KB，VmSwap升到1835092KB。这个结果非常直接地说明，当B的大工作集进入内存时，内核把A进程暂时不活跃的匿名页批量换出，两个进程之间确实发生了物理页竞争。")

    p(cell, "1.5 访问局部性验证", "Title1")
    image(cell, "07_locality_output.png", "终端显示全空间循环与16页分组访问结果", max_width=6.2, folder=TERM)
    p(cell, "全空间逐页循环在4096MB、重复3遍的设置下耗时51.264秒，主缺页为265845次；16页分组访问在总页访问次数不变的条件下仅耗时3.652秒，主缺页只有11次。两者差异说明，页面置换的性能损耗不只由访问总量决定，更由下一次访问发生时该页是否仍处在工作集中决定。16个连续页为一组时，组内64KB工作集能够在短时间内保持驻留，因此重复访问几乎不需要从swap换回。")

    p(cell, "2. 结论性评价", "Title2")
    p(cell, "本次实验最终证明：Linux进程的虚拟地址空间采用以VMA为单位的离散式管理，大块malloc会体现为匿名映射区的建立、拆分与合并；虚拟内存申请成功并不等价于物理内存立刻占用，读匿名页和写匿名页在缺页处理路径上存在本质差异；当工作集超过物理内存时，页面置换会把不活跃匿名页推入swap，而良好的局部性可以在访问总量不变的情况下显著减少主缺页和运行时间。")

    p(cell, "五、实验体会：（根据自己情况填写）", "Normal")
    p(cell, "最开始做这个实验时，我对“分配了内存”这句话的理解仍然有些粗糙，潜意识里总觉得malloc或mmap成功后，内核就已经把对应大小的物理内存交给了进程。真正跑完3GB读触页后，这个认知被数据直接推翻了：VmSize已经扩张到3GB级别，但RssAnon几乎没有变化。这个结果让我重新把虚拟地址、页表项和物理页三者拆开理解，也意识到操作系统课堂上讲的按需分页并不是抽象概念，而是可以在/proc文件中被逐项看见的真实机制。")
    p(cell, "调试过程中最有代表性的坑出现在双进程竞争实验里。我一开始准备用A进程自己的输出日志判断竞争前后的RSS变化，但A进程被重定向到文件后stdout进入全缓冲状态，中间状态并不能及时落盘。后来我改为在B进程运行前后由外部直接读取/proc/A/status，这才拿到A的VmRSS从约1.79GB跌到约1MB、VmSwap升到约1.75GB的关键证据。这个小问题提醒我，系统实验不能只相信程序自己的叙述，应该尽量从内核暴露的状态接口交叉验证。")
    p(cell, "局部性对比实验给我的冲击也很强。同样是4096MB空间、同样重复访问3遍，只是把访问顺序从全空间循环改成16页一组，时间就从51秒降到3秒多。这个差异让我更具体地理解了为什么操作系统教材会如此强调工作集模型：性能问题并不总是来自“做了多少工作”，也可能来自“工作以什么顺序抵达内存系统”。通过这次实验，我对虚拟内存、物理页、swap和页面置换之间的关系有了更底层也更工程化的认识。")


def main():
    doc = Document(BASE)

    table = doc.tables[0]
    for row_idx in [0, 1, 2]:
        clear_cell(table.rows[row_idx].cells[0])

    objective_num_id = create_paren_numbering(doc)
    constraint_num_id = create_paren_numbering(doc)
    compare_num_id = create_paren_numbering(doc)

    add_objectives(table.rows[0].cells[0], objective_num_id)
    add_methods(table.rows[1].cells[0], constraint_num_id)
    add_process(table.rows[1].cells[0], compare_num_id)
    add_results(table.rows[2].cells[0])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
